# -*- coding: utf-8 -*-
"""
Created on Wed Apr 22 20:42:37 2026

@author: jesus
"""
# Librería para comprobar si existe la plantilla Word
import os

# Librería para abrir y editar documentos Word
from docx import Document

# Importamos rutas del proyecto
from common.paths import ASSETS_DIR, OUTPUT_DIR

# Importamos funciones auxiliares reutilizables
from common.utils import set_cell_text, formatear_numero, formatear_calificacion


# =====================================================
# RUTA DE LA PLANTILLA WORD DE 105/51
# =====================================================
WORD_TEMPLATE_105_51 = os.path.join(ASSETS_DIR, "tabla_105_51.docx")


# =====================================================
# GENERAR WORD DE 105/51 A PARTIR DE RESULTADOS
# =====================================================
def generar_word_105_51(data: dict) -> str:
    """
    Genera el documento Word de munición 105/51
    a partir del diccionario de resultados calculados.

    Devuelve la ruta completa del archivo generado.
    """

    # =====================================================
    # COMPROBAR QUE EXISTE LA PLANTILLA WORD
    # =====================================================
    if not os.path.exists(WORD_TEMPLATE_105_51):
        raise FileNotFoundError(
            f"No existe la plantilla Word:\n{WORD_TEMPLATE_105_51}"
        )

    # =====================================================
    # ABRIR PLANTILLA Y SELECCIONAR LA PRIMERA TABLA
    # =====================================================
    doc = Document(WORD_TEMPLATE_105_51)

    if not doc.tables:
        raise ValueError("La plantilla Word no contiene ninguna tabla.")

    table = doc.tables[0]

    # =====================================================
    # FILA VELOCIDAD
    # =====================================================
    # ÚTIL-1: rango fijo
    # INÚTIL: fuera del rango útil-2
    # ÚTIL-2: resto de casos
    set_cell_text(
        table.cell(2, 1),
        f"V0c ∈ [{formatear_numero(data['r1_min'])}, {formatear_numero(data['r1_max'])}] m/s"
    )
    set_cell_text(
        table.cell(2, 2),
        f"V0c ∉ [{formatear_numero(data['r2_min'])}, {formatear_numero(data['r2_max'])}] m/s"
    )
    set_cell_text(
        table.cell(2, 3),
        "Resto de casos"
    )
    set_cell_text(
        table.cell(2, 4),
        f"V̄0c = {formatear_numero(data['v_med'])} m/s"
    )
    set_cell_text(table.cell(2, 5), formatear_calificacion(data["res_vel"]))

    # =====================================================
    # FILA DESVIACIÓN
    # =====================================================
    # =====================================================
# FILA DESVIACIÓN
# =====================================================
    set_cell_text(
    table.cell(3, 1),
    f"σn-1(V0) * {formatear_numero(data['desv_factor'])} ≤ {formatear_numero(data['desv_lim'])} m/s"
    )
    set_cell_text(
    table.cell(3, 2),
    f"σn-1(V0) * {formatear_numero(data['desv_factor'])} > {formatear_numero(data['desv_lim'])} m/s"
    )
    set_cell_text(
    table.cell(3, 3),
    "Resto de casos"
    )
    set_cell_text(
    table.cell(3, 4),
    f"σn-1(V0) * {formatear_numero(data['desv_factor'])} = {formatear_numero(data['desv_corregida'])} m/s"
    )
    set_cell_text(table.cell(3, 5), formatear_calificacion(data["res_desv"]))

    # =====================================================
    # FILA PMÁX
    # =====================================================
    set_cell_text(
        table.cell(4, 1),
        f"Pmáx ≤ {formatear_numero(data['pmax_lim'])} MPa"
    )
    set_cell_text(
        table.cell(4, 2),
        f"Pmáx > {formatear_numero(data['pmax_lim'])} MPa"
    )
    set_cell_text(
        table.cell(4, 3),
        "-"
    )
    set_cell_text(
        table.cell(4, 4),
        f"Pmax = {formatear_numero(data['pmax'])} MPa"
    )
    set_cell_text(table.cell(4, 5), formatear_calificacion(data["res_pmax"]))


    # =====================================================
    # FILA PMED
    # =====================================================
    set_cell_text(
        table.cell(5, 1),
        f"Pmed ≤ {formatear_numero(data['pmed_lim'])} MPa"
    )
    set_cell_text(
        table.cell(5, 2),
        f"Pmed > {formatear_numero(data['pmed_lim'])} MPa"
    )
    set_cell_text(
        table.cell(5, 3),
        "-"
    )
    set_cell_text(
        table.cell(5, 4),
        f"Pmed = {formatear_numero(data['pmed'])} MPa"
    )
    set_cell_text(table.cell(5, 5), formatear_calificacion(data["res_pmed"]))

    # =====================================================
    # FILA PMED + 3 * SIGMA(PMED)
    # =====================================================
    set_cell_text(
        table.cell(6, 1),
        f"Pmed + 3·σPmed ≤ {formatear_numero(data['pmed_sigma_lim'])} MPa"
    )
    set_cell_text(
        table.cell(6, 2),
        f"Pmed + 3·σPmed > {formatear_numero(data['pmed_sigma_lim'])} MPa"
    )
    set_cell_text(
        table.cell(6, 3),
        "-"
    )
    set_cell_text(
        table.cell(6, 4),
        f"Pmed + 3·σPmed = {formatear_numero(data['pmed_sigma'])} MPa"
    )
    set_cell_text(table.cell(6, 5), formatear_calificacion(data["res_pmed_sigma"]))

    # =====================================================
    # FILA PROYECTIL - SEPARACIÓN DE PARTES METÁLICAS
    # =====================================================
    # En la tabla la lógica útil-1 / útil-2 puede estar vacía y solo se
    # refleja el número de fallos y la calificación final.
    set_cell_text(table.cell(7, 4), f"{data['fallos_separacion']} fallos")
    set_cell_text(table.cell(7, 5), formatear_calificacion(data["res_separacion"]))

    # =====================================================
    # FILA PROYECTIL - ANOMALÍAS EN VUELO
    # =====================================================
    set_cell_text(table.cell(8, 4), f"{data['fallos_vuelo']} fallos")
    set_cell_text(table.cell(8, 5), formatear_calificacion(data["res_vuelo"]))

    # =====================================================
    # FILA ESPOLETA
    # =====================================================
    set_cell_text(table.cell(9, 1), "Fallos ≤ 2")
    set_cell_text(table.cell(9, 2), "Fallos > 2")
    set_cell_text(table.cell(9, 3), "-")
    set_cell_text(table.cell(9, 4), f"{data['fallos_espoleta']} fallos")
    set_cell_text(table.cell(9, 5), formatear_calificacion(data["res_espoleta"]))

    # =====================================================
    # FILA ESTOPÍN
    # =====================================================
    set_cell_text(table.cell(10, 1), "Fallos ≤ 1")
    set_cell_text(table.cell(10, 2), "Fallos > 1")
    set_cell_text(table.cell(10, 3), "-")
    set_cell_text(table.cell(10, 4), f"{data['fallos_estopin']} fallos")
    set_cell_text(table.cell(10, 5), formatear_calificacion(data["res_estopin"]))

    # =====================================================
    # FILA FUGA DE GASES POR ESTOPÍN
    # =====================================================
    # NO SE RELLENA: contenido estático en la plantilla Word

    # =====================================================
    # FILA ROTURA CÁPSULA ESTOPÍN
    # =====================================================
    # NO SE RELLENA: contenido estático en la plantilla Word

    # =====================================================
    # FILA TRAZADOR
    # =====================================================
    set_cell_text(table.cell(13, 1), "-")
    set_cell_text(table.cell(13, 2), "3 trazador < 2,15 s")
    set_cell_text(table.cell(13, 3), "-")
    set_cell_text(
        table.cell(13, 4),
        f"{data['vuelos_trazador_fuera']} vuelos tienen un tiempo de trazador menor de {formatear_numero(data['trazador_lim'])} s"
    )
    set_cell_text(table.cell(13, 5), formatear_calificacion(data["res_trazador"]))

    # =====================================================
    # GUARDAR DOCUMENTO EN OUTPUT
    # =====================================================
    nombre_archivo = f"{data['nombre_prueba']}.docx"
    out_path = os.path.join(OUTPUT_DIR, nombre_archivo)
    doc.save(out_path)

    return out_path
