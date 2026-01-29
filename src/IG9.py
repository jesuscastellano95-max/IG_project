# -*- coding: utf-8 -*-
"""
Created on Sat Dec 20 08:29:57 2025

@author: jesus
"""
import os
import tkinter as tk
from tkinter import ttk, messagebox
from PIL import Image, ImageTk
import pandas as pd
from docx import Document
from docx.shared import Pt

# =====================================================
# RUTAS (estructura del proyecto)
# IG_project/
#   src/IG8.py
#   assets/ (logo + plantilla Word)
#   data/   (excel)
#   output/ (word generados)
# =====================================================
SRC_DIR = os.path.dirname(os.path.abspath(__file__))              # .../IG_project/src
PROJECT_DIR = os.path.abspath(os.path.join(SRC_DIR, ".."))       # .../IG_project

ASSETS_DIR = os.path.join(PROJECT_DIR, "assets")
DATA_DIR = os.path.join(PROJECT_DIR, "data")
OUTPUT_DIR = os.path.join(PROJECT_DIR, "output")

EXCEL_PATH = os.path.join(DATA_DIR, "excelplantilla.xlsx")
WORD_TEMPLATE_PATH = os.path.join(ASSETS_DIR, "cuadro_resultados_general.docx")
LOGO_PATH = os.path.join(ASSETS_DIR, "SIIGROUPLOGO.jpg")

# Asegurar que existe la carpeta de salida
os.makedirs(OUTPUT_DIR, exist_ok=True)

# =====================================================
# CONFIG ARMAS (preparado para crecer)
# =====================================================
ARMAS = {
    "OTO MELARA 76/62": {
        "sheet": "OTO_MELARA",
        "vnom_fields": ["Vnom (m/s)"],
        "pmed_lim": 3285.23,
        "pmax_lim": 3522.55,
        "desv_lim": 5.0,
    },
    "Granadas de mortero rompedoras y fumígenas": {
        "sheet": "MORTERO",
        "vnom_fields": ["Vnom carga 1 (m/s)", "Vnom carga máxima (m/s)"],
        # límites / constantes del mortero las meteremos más adelante
    }
}

# =====================================================
# FUNCIONES AUXILIARES (IG7)
# =====================================================

def convertir_a_float(texto: str) -> float:
    """Acepta decimales con coma o punto."""
    return float(texto.strip().replace(",", "."))

def formatear_numero(valor: float) -> str:
    """
    - Si es entero: sin decimales
    - Si es decimal: 2 decimales y con coma
    """
    try:
        v = float(valor)
    except Exception:
        return str(valor)

    if v.is_integer():
        return f"{int(v)}"
    return f"{v:.2f}".replace(".", ",")

def set_cell_text(cell, text: str):
    """Escribe en Times New Roman 10."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

def serie_numerica(df: pd.DataFrame, col: str) -> pd.Series:
    """
    Convierte una columna a numérica aunque venga como texto con comas decimales.
    """
    s = df[col]
    if s.dtype == object:
        s = s.astype(str).str.replace(",", ".", regex=False)
    s = pd.to_numeric(s, errors="coerce")
    return s.dropna()

def contar_fallos(df: pd.DataFrame, col: str) -> int:
    """
    Cuenta cuántas veces aparece 'Fallo' (ignora mayúsculas/minúsculas y espacios).
    """
    s = df[col].astype(str).str.strip().str.lower()
    return int((s == "fallo").sum())

# =====================================================
# ESTADO (para pasar de Pantalla 2 -> Pantalla 3)
# =====================================================
RESULTADOS = None  # dict con valores y calificaciones
ARMA_SELECCIONADA = None
VNOMS_GLOBAL = {}  # para armas con 1 o 2 velocidades nominales


# =====================================================
# CÁLCULO (MISMA LÓGICA QUE IG7) — NO genera Word
# =====================================================

def calcular_resultados():
    """
    Calcula TODO igual que IG7 (misma lógica),
    pero devuelve un dict con valores y resultados.
    """
    # -------- INPUTS MANUALES --------
    nombre_prueba = entry_nombre_prueba.get().strip()
    if not nombre_prueba:
        raise ValueError("Introduce un nombre de prueba (será el nombre del Word).")

    r1_min = convertir_a_float(entry_r1_min.get())
    r1_max = convertir_a_float(entry_r1_max.get())
    r2_min = convertir_a_float(entry_r2_min.get())
    r2_max = convertir_a_float(entry_r2_max.get())

    limite_desv = convertir_a_float(entry_desv_lim.get())
    pmed_lim = convertir_a_float(entry_pmed_lim.get())
    pmax_lim = convertir_a_float(entry_pmax_lim.get())

    # -------- CARGA EXCEL --------
    if not os.path.exists(EXCEL_PATH):
        raise FileNotFoundError(f"No existe el archivo:\n{EXCEL_PATH}")

    df = pd.read_excel(EXCEL_PATH)

    # -------- COLUMNAS EXCEL --------
    col_vel = entry_col_vel.get().strip()
    col_desv = entry_col_desv.get().strip()
    col_pmed = entry_col_pmed.get().strip()
    col_pmax1 = entry_col_pmax1.get().strip()
    col_pmax2 = entry_col_pmax2.get().strip()

    col_espoleta = entry_col_espoleta.get().strip()
    col_estopin = entry_col_estopin.get().strip()

    # Comprobación de existencia de columnas
    columnas_necesarias = [col_vel, col_desv, col_pmed, col_pmax1, col_pmax2, col_espoleta, col_estopin]
    for c in columnas_necesarias:
        if not c:
            raise ValueError("Hay columnas del Excel sin rellenar en la interfaz.")
        if c not in df.columns:
            raise ValueError(f"La columna '{c}' no existe en el Excel.")

    # -------- CÁLCULOS DESDE EXCEL --------
    v_med = serie_numerica(df, col_vel).mean()
    desv = serie_numerica(df, col_desv).std()
    pmed = serie_numerica(df, col_pmed).mean()

    pmax1 = serie_numerica(df, col_pmax1).max()
    pmax2 = serie_numerica(df, col_pmax2).max()
    pmax = max(pmax1, pmax2)

    fallos_espoleta = contar_fallos(df, col_espoleta)
    fallos_estopin = contar_fallos(df, col_estopin)

    # -------- LÓGICA --------
    # Velocidad: UTIL-1 dentro rango estrecho; UTIL-2 dentro rango ancho; INUTIL fuera de ambos
    if r1_min <= v_med <= r1_max:
        res_vel = "UTIL-1"
    elif r2_min <= v_med <= r2_max:
        res_vel = "UTIL-2"
    else:
        res_vel = "INUTIL"

    # Desviación estándar: UTIL-1 si menor que límite, si no INUTIL
    res_desv = "UTIL-1" if desv < limite_desv else "INUTIL"

    # Pmáx:
    # INUTIL solo si ALGUNA presión individual > límite. Si ambas <=, entonces UTIL-1
    if (pmax1 > pmax_lim) or (pmax2 > pmax_lim):
        res_pmax = "INUTIL"
    else:
        res_pmax = "UTIL-1"

    # Presión media: UTIL-1 si menor que límite, si no INUTIL
    res_pmed = "UTIL-1" if pmed < pmed_lim else "INUTIL"

    # Espoleta: si fallos >= 2 => INUTIL, si no => UTIL-1
    res_espoleta = "INUTIL" if fallos_espoleta >= 2 else "UTIL-1"

    # Estopín: si fallos >= 1 => INUTIL, si no => UTIL-1
    res_estopin = "INUTIL" if fallos_estopin >= 1 else "UTIL-1"

    return {
        "nombre_prueba": nombre_prueba,
        "r1_min": r1_min, "r1_max": r1_max,
        "r2_min": r2_min, "r2_max": r2_max,
        "limite_desv": limite_desv,
        "pmed_lim": pmed_lim,
        "pmax_lim": pmax_lim,
        "v_med": v_med,
        "desv": desv,
        "pmed": pmed,
        "pmax": pmax,
        "pmax1": pmax1,
        "pmax2": pmax2,
        "fallos_espoleta": fallos_espoleta,
        "fallos_estopin": fallos_estopin,
        "res_vel": res_vel,
        "res_desv": res_desv,
        "res_pmax": res_pmax,
        "res_pmed": res_pmed,
        "res_espoleta": res_espoleta,
        "res_estopin": res_estopin,
    }

# =====================================================
# WORD (MISMO RELLENADO QUE IG7) — usa RESULTADOS
# =====================================================

def generar_word_desde_resultados(data: dict):
    if not os.path.exists(WORD_TEMPLATE_PATH):
        raise FileNotFoundError(f"No existe la plantilla Word:\n{WORD_TEMPLATE_PATH}")

    doc = Document(WORD_TEMPLATE_PATH)
    table = doc.tables[0]

    # Fila velocidad (row=2)
    set_cell_text(table.cell(2, 1), f"V0c ∈ [{formatear_numero(data['r1_min'])}, {formatear_numero(data['r1_max'])}] m/s")
    set_cell_text(table.cell(2, 2), f"V0c ∉ [{formatear_numero(data['r2_min'])}, {formatear_numero(data['r2_max'])}] m/s")
    set_cell_text(table.cell(2, 4), f"V̄0c = {formatear_numero(data['v_med'])} m/s")
    set_cell_text(table.cell(2, 5), data["res_vel"])

    # Fila desviación estándar (row=3)
    set_cell_text(table.cell(3, 1), f"σV0 ≤ {formatear_numero(data['limite_desv'])} m/s")
    set_cell_text(table.cell(3, 2), f"σV0 > {formatear_numero(data['limite_desv'])} m/s")
    set_cell_text(table.cell(3, 4), f"σV0 = {formatear_numero(data['desv'])} m/s")
    set_cell_text(table.cell(3, 5), data["res_desv"])

    # Fila Pmáx (row=4)
    set_cell_text(table.cell(4, 1), f"Pmáx ≤ {formatear_numero(data['pmax_lim'])} bar")
    set_cell_text(table.cell(4, 2), f"Pmáx > {formatear_numero(data['pmax_lim'])} bar")
    set_cell_text(table.cell(4, 4), f"Pmax = {formatear_numero(data['pmax'])} bar")
    set_cell_text(table.cell(4, 5), data["res_pmax"])

    # Fila Pmed (row=5)
    set_cell_text(table.cell(5, 1), f"Pmed < {formatear_numero(data['pmed_lim'])} bar")
    set_cell_text(table.cell(5, 2), f"Pmed ≥ {formatear_numero(data['pmed_lim'])} bar")
    set_cell_text(table.cell(5, 4), f"Pmed = {formatear_numero(data['pmed'])} bar")
    set_cell_text(table.cell(5, 5), data["res_pmed"])

    # Fila Espoleta (row=6)
    set_cell_text(table.cell(6, 4), f"{data['fallos_espoleta']} fallos")
    set_cell_text(table.cell(6, 5), data["res_espoleta"])

    # Fila Estopín (row=7)
    set_cell_text(table.cell(7, 4), f"{data['fallos_estopin']} fallos")
    set_cell_text(table.cell(7, 5), data["res_estopin"])

    out_path = os.path.join(OUTPUT_DIR, f"{data['nombre_prueba']}.docx")
    doc.save(out_path)
    return out_path

# =====================================================
# PANTALLA 3 — Confirmación y generar Word
# =====================================================

def abrir_pantalla_3():
    global ventana3

    if RESULTADOS is None:
        messagebox.showerror("Error", "Primero debes calcular los resultados.")
        return

    ventana3 = tk.Toplevel(ventana2)
    ventana3.title("IG8 remasterizado - Confirmación")

    # Header
    header = ttk.Frame(ventana3)
    header.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(10, 5), padx=10)

    try:
        img = Image.open(LOGO_PATH).resize((250, 100))
        ventana3.logo_img = ImageTk.PhotoImage(img)
        ttk.Label(header, image=ventana3.logo_img).grid(row=0, column=0, rowspan=2, padx=(0, 12))
    except Exception:
        ttk.Label(header, text="(Logo no encontrado)").grid(row=0, column=0, rowspan=2, padx=(0, 12))

    ttk.Label(header, text=f"Arma: {ARMA_SELECCIONADA}", font=("Arial", 12, "bold")).grid(row=0, column=1, sticky="w")
    ttk.Label(header, text=f"Velocidad nominal (Vnom): {int(round(VNOM_GLOBAL))} m/s", font=("Arial", 11)).grid(row=1, column=1, sticky="w")

    # Resultados (tabla visual)
    frame = ttk.LabelFrame(ventana3, text="Resultados calculados (confirmación)")
    frame.grid(row=1, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")

    filas = [
        ("Velocidad media", f"{formatear_numero(RESULTADOS['v_med'])} m/s", RESULTADOS["res_vel"]),
        ("Desv. estándar", f"{formatear_numero(RESULTADOS['desv'])} m/s", RESULTADOS["res_desv"]),
        ("Pmáx", f"{formatear_numero(RESULTADOS['pmax'])} bar", RESULTADOS["res_pmax"]),
        ("Pmed", f"{formatear_numero(RESULTADOS['pmed'])} bar", RESULTADOS["res_pmed"]),
        ("Espoleta", f"{RESULTADOS['fallos_espoleta']} fallos", RESULTADOS["res_espoleta"]),
        ("Estopín", f"{RESULTADOS['fallos_estopin']} fallos", RESULTADOS["res_estopin"]),
    ]

    ttk.Label(frame, text="Condición", font=("Arial", 10, "bold")).grid(row=0, column=0, padx=6, pady=4, sticky="w")
    ttk.Label(frame, text="Valor", font=("Arial", 10, "bold")).grid(row=0, column=1, padx=6, pady=4, sticky="w")
    ttk.Label(frame, text="Calificación", font=("Arial", 10, "bold")).grid(row=0, column=2, padx=6, pady=4, sticky="w")

    for i, (nom, val, cal) in enumerate(filas, start=1):
        ttk.Label(frame, text=nom).grid(row=i, column=0, padx=6, pady=3, sticky="w")
        ttk.Label(frame, text=val).grid(row=i, column=1, padx=6, pady=3, sticky="w")
        ttk.Label(frame, text=cal, font=("Arial", 10, "bold")).grid(row=i, column=2, padx=6, pady=3, sticky="w")

    # Botones
    btns = ttk.Frame(ventana3)
    btns.grid(row=2, column=0, columnspan=2, pady=10)

    def volver_a_pantalla_2():
        ventana3.destroy()
        ventana2.deiconify()

    def generar_word():
        try:
            out_path = generar_word_desde_resultados(RESULTADOS)
            messagebox.showinfo("IG8 remasterizado", f"Documento Word generado correctamente:\n{out_path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    ttk.Button(btns, text="Volver", command=volver_a_pantalla_2).grid(row=0, column=0, padx=6)
    ttk.Button(btns, text="Generar Word", command=generar_word).grid(row=0, column=1, padx=6)

    # Footer
    ttk.Label(
        ventana3,
        text="Software diseñado y creado por Jesús Castellano Chillas",
        font=("Arial", 9),
        foreground="gray"
    ).grid(row=3, column=0, columnspan=2, pady=(0, 10))

    # Ocultamos pantalla 2 mientras confirmamos
    ventana2.withdraw()

# =====================================================
# PANTALLA 2 (IG7) — NO TOCAR LÓGICA / WORD
# Cambios: botón pasa a "Calcular" + botón volver pantalla 1
# =====================================================

def abrir_pantalla_2(arma: str, vnom: float):
    global ventana2, ARMA_SELECCIONADA, VNOM_GLOBAL
    ARMA_SELECCIONADA = arma
    VNOM_GLOBAL = vnom

    ventana2 = tk.Toplevel(ventana1)
    ventana2.title("IG8 remasterizado - Evaluación Automática")

    # --- Header (logo + resumen) ---
    header = ttk.Frame(ventana2)
    header.grid(row=0, column=0, columnspan=4, sticky="ew", pady=(10, 5), padx=10)

    try:
        img = Image.open(LOGO_PATH).resize((250, 100))
        ventana2.logo_img = ImageTk.PhotoImage(img)
        ttk.Label(header, image=ventana2.logo_img).grid(row=0, column=0, rowspan=2, padx=(0, 12))
    except Exception:
        ttk.Label(header, text="(Logo no encontrado)").grid(row=0, column=0, rowspan=2, padx=(0, 12))

    ttk.Label(header, text=f"Arma: {arma}", font=("Arial", 12, "bold")).grid(row=0, column=1, sticky="w")
    ttk.Label(header, text=f"Velocidad nominal (Vnom): {int(round(vnom))} m/s", font=("Arial", 11)).grid(row=1, column=1, sticky="w")

    # --- Contenedor principal ---
    main = ttk.Frame(ventana2)
    main.grid(row=1, column=0, columnspan=4, sticky="nsew", padx=10, pady=5)

    # 2 columnas visuales: Inputs manuales / Columnas Excel
    lf_inputs = ttk.LabelFrame(main, text="Inputs (manuales)")
    lf_inputs.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=5)

    lf_excel = ttk.LabelFrame(main, text="Columnas (Excel)")
    lf_excel.grid(row=0, column=1, sticky="nsew", pady=5)

    # --- Inputs manuales ---
    r = 0
    ttk.Label(lf_inputs, text="Nombre de la prueba").grid(row=r, column=0, sticky="w", padx=6, pady=4)
    global entry_nombre_prueba
    entry_nombre_prueba = ttk.Entry(lf_inputs, width=28)
    entry_nombre_prueba.grid(row=r, column=1, columnspan=2, sticky="w", padx=6, pady=4)

    r += 1
    ttk.Label(lf_inputs, text="Rango UTIL-1 (min / max)").grid(row=r, column=0, sticky="w", padx=6, pady=4)
    global entry_r1_min, entry_r1_max
    entry_r1_min = ttk.Entry(lf_inputs, width=10)
    entry_r1_min.grid(row=r, column=1, sticky="w", padx=6, pady=4)
    entry_r1_max = ttk.Entry(lf_inputs, width=10)
    entry_r1_max.grid(row=r, column=2, sticky="w", padx=6, pady=4)

    r += 1
    ttk.Label(lf_inputs, text="Rango UTIL-2 (min / max)").grid(row=r, column=0, sticky="w", padx=6, pady=4)
    global entry_r2_min, entry_r2_max
    entry_r2_min = ttk.Entry(lf_inputs, width=10)
    entry_r2_min.grid(row=r, column=1, sticky="w", padx=6, pady=4)
    entry_r2_max = ttk.Entry(lf_inputs, width=10)
    entry_r2_max.grid(row=r, column=2, sticky="w", padx=6, pady=4)

    r += 1
    ttk.Label(lf_inputs, text="Desv. estándar límite").grid(row=r, column=0, sticky="w", padx=6, pady=4)
    global entry_desv_lim
    entry_desv_lim = ttk.Entry(lf_inputs, width=12)
    entry_desv_lim.grid(row=r, column=1, sticky="w", padx=6, pady=4)

    r += 1
    ttk.Label(lf_inputs, text="Presión media límite").grid(row=r, column=0, sticky="w", padx=6, pady=4)
    global entry_pmed_lim
    entry_pmed_lim = ttk.Entry(lf_inputs, width=12)
    entry_pmed_lim.grid(row=r, column=1, sticky="w", padx=6, pady=4)

    r += 1
    ttk.Label(lf_inputs, text="Presión individual máx. permitida").grid(row=r, column=0, sticky="w", padx=6, pady=4)
    global entry_pmax_lim
    entry_pmax_lim = ttk.Entry(lf_inputs, width=12)
    entry_pmax_lim.grid(row=r, column=1, sticky="w", padx=6, pady=4)

    # --- Columnas Excel ---
    e = 0
    ttk.Label(lf_excel, text="Columna velocidad media").grid(row=e, column=0, sticky="w", padx=6, pady=4)
    global entry_col_vel
    entry_col_vel = ttk.Entry(lf_excel, width=28)
    entry_col_vel.grid(row=e, column=1, sticky="w", padx=6, pady=4)

    e += 1
    ttk.Label(lf_excel, text="Columna desviación estándar").grid(row=e, column=0, sticky="w", padx=6, pady=4)
    global entry_col_desv
    entry_col_desv = ttk.Entry(lf_excel, width=28)
    entry_col_desv.grid(row=e, column=1, sticky="w", padx=6, pady=4)

    e += 1
    ttk.Label(lf_excel, text="Columna presión media").grid(row=e, column=0, sticky="w", padx=6, pady=4)
    global entry_col_pmed
    entry_col_pmed = ttk.Entry(lf_excel, width=28)
    entry_col_pmed.grid(row=e, column=1, sticky="w", padx=6, pady=4)

    e += 1
    ttk.Label(lf_excel, text="Columnas Pmáx (2 columnas)").grid(row=e, column=0, sticky="w", padx=6, pady=4)
    global entry_col_pmax1, entry_col_pmax2
    frame_pmax = ttk.Frame(lf_excel)
    frame_pmax.grid(row=e, column=1, sticky="w", padx=6, pady=4)
    entry_col_pmax1 = ttk.Entry(frame_pmax, width=13)
    entry_col_pmax1.grid(row=0, column=0, padx=(0, 6))
    entry_col_pmax2 = ttk.Entry(frame_pmax, width=13)
    entry_col_pmax2.grid(row=0, column=1)

    e += 1
    ttk.Label(lf_excel, text="Columna Espoleta  (<2 fallos)").grid(row=e, column=0, sticky="w", padx=6, pady=4)
    global entry_col_espoleta
    entry_col_espoleta = ttk.Entry(lf_excel, width=28)
    entry_col_espoleta.grid(row=e, column=1, sticky="w", padx=6, pady=4)

    e += 1
    ttk.Label(lf_excel, text="Columna Estopín  (<1 fallos)").grid(row=e, column=0, sticky="w", padx=6, pady=4)
    global entry_col_estopin
    entry_col_estopin = ttk.Entry(lf_excel, width=28)
    entry_col_estopin.grid(row=e, column=1, sticky="w", padx=6, pady=4)

    # --- Botones (Calcular + Volver a pantalla 1) ---
    btn_frame = ttk.Frame(ventana2)
    btn_frame.grid(row=2, column=0, columnspan=4, pady=12)

    def accion_calcular():
        global RESULTADOS
        try:
            RESULTADOS = calcular_resultados()
        except Exception as e2:
            messagebox.showerror("Error", str(e2))
            return
        abrir_pantalla_3()

    def volver_a_pantalla_1():
        # cerrar pantalla 2 y volver a 1
        ventana2.destroy()
        ventana1.deiconify()

    ttk.Button(btn_frame, text="Calcular", command=accion_calcular).grid(row=0, column=0, padx=6)
    ttk.Button(btn_frame, text="Volver a Pantalla 1", command=volver_a_pantalla_1).grid(row=0, column=1, padx=6)

    # --- Footer ---
    footer = ttk.Label(
        ventana2,
        text="Software diseñado y creado por Jesús Castellano Chillas",
        font=("Arial", 9),
        foreground="gray"
    )
    footer.grid(row=3, column=0, columnspan=4, pady=(0, 10))

    # --- Autocompletar campos desde Pantalla 1 (EDITABLES) ---
    cfg = ARMAS.get(arma, {})

    # Rangos como ENTEROS
    r1_min = int(round(vnom - 7))
    r1_max = int(round(vnom + 7))
    r2_min = int(round(vnom - 0.02 * vnom))
    r2_max = int(round(vnom + 0.02 * vnom))

    entry_r1_min.insert(0, str(r1_min))
    entry_r1_max.insert(0, str(r1_max))
    entry_r2_min.insert(0, str(r2_min))
    entry_r2_max.insert(0, str(r2_max))

    # Fijar límites de arma (OTO MELARA)
    if "desv_lim" in cfg:
        entry_desv_lim.insert(0, formatear_numero(cfg["desv_lim"]))  # 5
    if "pmed_lim" in cfg:
        entry_pmed_lim.insert(0, formatear_numero(cfg["pmed_lim"]))  # 3285,23
    if "pmax_lim" in cfg:
        entry_pmax_lim.insert(0, formatear_numero(cfg["pmax_lim"]))  # 3522,55

    # Cerrar/ocultar pantalla 1 (obligatoria)
    ventana1.withdraw()

# =====================================================
# PANTALLA 1 (OBLIGATORIA)
# =====================================================

def continuar_a_pantalla_2():
    global VNOMS_GLOBAL

    arma = combo_arma.get().strip()
    if arma not in ARMAS:
        messagebox.showerror("Error", "Selecciona un arma válida.")
        return

    cfg = ARMAS.get(arma, {})
    campos = cfg.get("vnom_fields", ["Vnom (m/s)"])

    vnoms = {}
    try:
        for c in campos:
            texto = vnom_entries[c].get()
            vnoms[c] = convertir_a_float(texto)
    except Exception:
        messagebox.showerror("Error", "Introduce velocidades nominales válidas (acepta coma o punto).")
        return

    VNOMS_GLOBAL = vnoms

    # Para OTO MELARA (1 velocidad), comportamiento idéntico al actual:
    if len(vnoms) == 1:
        vnom_unico = list(vnoms.values())[0]
        abrir_pantalla_2(arma, vnom_unico)
    else:
        # Mortero: por ahora solo pasamos una (para no tocar Pantalla 2 todavía)
        vnom_cmax = vnoms.get("Vnom carga máxima (m/s)", list(vnoms.values())[0])
        abrir_pantalla_2(arma, vnom_cmax)

ventana1 = tk.Tk()
ventana1.title("IG8 remasterizado - Inicio")
ventana1.geometry("700x320")   # ancho x alto (puedes variar)
ventana1.resizable(False, False)  # opcional: evita que se redimensione

# Logo
try:
    img = Image.open(LOGO_PATH).resize((250, 100))
    ventana1.logo_img = ImageTk.PhotoImage(img)
    ttk.Label(ventana1, image=ventana1.logo_img).grid(row=0, column=0, columnspan=2, pady=10, padx=10)
except Exception:
    ttk.Label(ventana1, text="(Logo no encontrado)").grid(row=0, column=0, columnspan=2, pady=10)

# Controles
ttk.Label(ventana1, text="Arma").grid(row=1, column=0, sticky="w", padx=10, pady=6)
combo_arma = ttk.Combobox(ventana1, values=list(ARMAS.keys()), state="readonly", width=45)
combo_arma.grid(row=1, column=1, sticky="w", padx=10, pady=6)
combo_arma.set("OTO MELARA 76/62")

# Velocidades nominales (dinámicas)
ttk.Label(ventana1, text="Velocidades nominales (m/s)").grid(row=2, column=0, sticky="w", padx=10, pady=6)

frame_vnom = ttk.Frame(ventana1)
frame_vnom.grid(row=2, column=1, sticky="w", padx=10, pady=6)

vnom_entries = {}  # guardaremos aquí las cajas de texto

def actualizar_campos_vnom(*args):
    # borrar widgets anteriores
    for w in frame_vnom.winfo_children():
        w.destroy()
    vnom_entries.clear()

    arma_sel = combo_arma.get().strip()
    cfg = ARMAS.get(arma_sel, {})
    campos = cfg.get("vnom_fields", ["Vnom (m/s)"])

    for i, label in enumerate(campos):
        ttk.Label(frame_vnom, text=label).grid(row=i, column=0, sticky="w", pady=2)
        ent = ttk.Entry(frame_vnom, width=12)
        ent.grid(row=i, column=1, sticky="w", padx=(6, 0), pady=2)
        vnom_entries[label] = ent

combo_arma.bind("<<ComboboxSelected>>", actualizar_campos_vnom)
actualizar_campos_vnom()  # crea los campos al arrancar

#Boton continuar
ttk.Button(ventana1, text="Continuar", command=continuar_a_pantalla_2).grid(
    row=3, column=0, columnspan=2, pady=12
)

# Footer
ttk.Label(
    ventana1,
    text="Software diseñado y creado por Jesús Castellano Chillas",
    font=("Arial", 9),
    foreground="gray"
).grid(row=4, column=0, columnspan=2, pady=(0, 10))

ventana1.mainloop()

